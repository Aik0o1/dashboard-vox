# relatorio.py
from docx import Document
from docx.shared import Inches
import plotly.io as pio

import plotly.io as pio

def salvar_grafico_png(figura, nome_arquivo):
    # Salva a figura como um arquivo PNG
    pio.write_image(figura, f"./relatorio/imgs/{nome_arquivo}.png", format="png")


def gerar_relatorio_word(graficos, nomes_arquivos, nome_relatorio):
    # Cria um novo documento Word
    doc = Document()

    # Adiciona um título ao documento
    doc.add_heading("Relatório de Gráficos", level=1)

    # Adiciona cada gráfico ao documento
    for grafico, nome_arquivo in zip(graficos, nomes_arquivos):
        # Salva o gráfico como SVG
        salvar_grafico_png(grafico, nome_arquivo)

        # Adiciona o gráfico ao documento Word
        doc.add_paragraph(nome_arquivo)
        doc.add_picture(f"./relatorio/imgs/{nome_arquivo}.png", width=Inches(6))

    # Salva o documento Word
    doc.save(f"{nome_relatorio}.docx")
